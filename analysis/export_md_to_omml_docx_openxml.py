from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt


PAPER_DIR = Path(r"F:\IJDRR write\论文初稿模块")
TARGETS = [
    PAPER_DIR / "IJDRR_main_manuscript_v2_refs_resolved.md",
    PAPER_DIR / "IJDRR_full_paper_v2_refs_resolved.md",
]

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
INLINE_MATH_RE = re.compile(r"\\\((.+?)\\\)")
IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def strip_caption_marks(text: str) -> tuple[str, bool]:
    stripped = text.strip()
    if len(stripped) >= 2 and stripped.startswith("*") and stripped.endswith("*"):
        return clean_text(stripped[1:-1]), True
    return clean_text(text), False


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_text(cell.strip()) for cell in line.split("|")]


def resolve_image(md_path: Path, raw_path: str) -> Path:
    raw_path = raw_path.replace("/", "\\")
    candidate = Path(raw_path)
    if candidate.is_absolute():
        return candidate
    return (md_path.parent / candidate).resolve()


def prepare_math(math: str) -> tuple[str, str | None]:
    math = math.strip()
    tag = None
    match = re.search(r"\\tag\{([^}]+)\}", math)
    if match:
        tag = match.group(1)
        math = math[: match.start()] + math[match.end() :]
    math = re.sub(r"\\begin\{aligned\}", "", math)
    math = re.sub(r"\\end\{aligned\}", "", math)
    math = re.sub(r"\\begin\{align\*?\}", "", math)
    math = re.sub(r"\\end\{align\*?\}", "", math)
    return math.strip(), tag


def make_omath(math: str):
    math, _ = prepare_math(math)
    return parse_xml(
        f'<m:oMath xmlns:m="{M_NS}"><m:r><m:t xml:space="preserve">{escape(math)}</m:t></m:r></m:oMath>'
    )


def make_omath_para(math: str):
    math, _ = prepare_math(math)
    return parse_xml(
        f'<m:oMathPara xmlns:m="{M_NS}"><m:oMath><m:r><m:t xml:space="preserve">{escape(math)}</m:t></m:r></m:oMath></m:oMathPara>'
    )


class DocxBuilder:
    def __init__(self, md_path: Path):
        self.md_path = md_path
        self.doc = Document()
        self.math_count = 0
        self.image_count = 0
        self.table_count = 0
        self._setup_styles()

    def _setup_styles(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(11)

    def add_heading(self, level: int, text: str) -> None:
        p = self.doc.add_heading(clean_text(text), level=min(level, 3))
        for run in p.runs:
            run.font.name = "Times New Roman"

    def add_paragraph(self, text: str, *, italic: bool = False, code: bool = False) -> None:
        p = self.doc.add_paragraph()
        if code:
            p.style = self.doc.styles["No Spacing"]
        last = 0
        for match in INLINE_MATH_RE.finditer(text):
            if match.start() > last:
                run = p.add_run(clean_text(text[last : match.start()]))
                run.italic = italic
                if code:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            run = p.add_run()
            run._r.append(make_omath(match.group(1)))
            self.math_count += 1
            last = match.end()
        if last < len(text):
            run = p.add_run(clean_text(text[last:]))
            run.italic = italic
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)

    def add_display_math(self, math: str) -> None:
        math, tag = prepare_math(math)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(make_omath_para(math))
        self.math_count += 1
        if tag:
            p.add_run(f"    ({tag})")

    def add_image(self, raw_path: str) -> None:
        image_path = resolve_image(self.md_path, raw_path)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if image_path.exists():
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(6.2))
            self.image_count += 1
        else:
            p.add_run(f"[Missing image: {raw_path}]")

    def add_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        cols = max(len(row) for row in rows)
        rows = [row + [""] * (cols - len(row)) for row in rows]
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(9)
                        if r_idx == 0:
                            run.bold = True
        self.table_count += 1

    def save(self, out_path: Path) -> None:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(out_path)


def parse_markdown(md_path: Path, builder: DocxBuilder) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            for code_line in code_lines:
                builder.add_paragraph(code_line, code=True)
            i += 1
            continue

        if stripped == "$$" or stripped.startswith("$$"):
            math_lines: list[str] = []
            if stripped != "$$":
                initial = stripped[2:]
                if initial.endswith("$$"):
                    builder.add_display_math(initial[:-2])
                    i += 1
                    continue
                math_lines.append(initial)
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                math_lines.append(lines[i])
                i += 1
            builder.add_display_math("\n".join(math_lines))
            i += 1
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            builder.add_image(image_match.group("path"))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1]):
            table_rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(split_table_row(lines[i]))
                i += 1
            builder.add_table(table_rows)
            continue

        if stripped.startswith("#"):
            match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if match:
                builder.add_heading(len(match.group(1)), match.group(2))
                i += 1
                continue

        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            if (
                nxt_stripped.startswith("#")
                or nxt_stripped.startswith("```")
                or nxt_stripped == "$$"
                or nxt_stripped.startswith("$$")
                or IMAGE_RE.match(nxt_stripped)
                or (nxt_stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1]))
            ):
                break
            para_lines.append(nxt_stripped)
            i += 1
        paragraph = " ".join(para_lines)
        paragraph, italic = strip_caption_marks(paragraph)
        builder.add_paragraph(paragraph, italic=italic)


def inspect_docx(path: Path) -> dict[str, int | bool]:
    with ZipFile(path) as zf:
        names = zf.namelist()
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return {
        "exists": path.exists(),
        "omath": xml.count("<m:oMath"),
        "paragraphs": xml.count("<w:p"),
        "tables": xml.count("<w:tbl"),
        "images": sum(1 for name in names if name.startswith("word/media/")),
    }


def convert_one(md_path: Path) -> tuple[Path, dict[str, int | bool], DocxBuilder]:
    builder = DocxBuilder(md_path)
    parse_markdown(md_path, builder)
    out_path = md_path.with_name(md_path.stem + "_OMML.docx")
    builder.save(out_path)
    return out_path, inspect_docx(out_path), builder


def main() -> int:
    for md_path in TARGETS:
        out_path, stats, builder = convert_one(md_path)
        print(
            f"{out_path.name}: omath={stats['omath']}, images={stats['images']}, "
            f"tables={stats['tables']}, paragraphs={stats['paragraphs']}, "
            f"builder_math={builder.math_count}"
        )
        print(f"  {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
