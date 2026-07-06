"""Convert project Markdown manuscripts to reviewer-friendly DOCX files.

This converter is intentionally small and local: it preserves manuscript
structure, Markdown tables, embedded figures and display-equation blocks well
enough for supervisor review without depending on pandoc.
"""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path
from urllib.parse import unquote

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
MARGIN_CM = 2.54
CONTENT_WIDTH_IN = (A4_WIDTH_CM - 2 * MARGIN_CM) / 2.54
CURRENT_INLINE_MATH_BLOCKS: list[str] = []


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "BFBFBF") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_horizontal_rule(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def preprocess_display_math(markdown_text: str) -> tuple[str, list[str], list[str]]:
    """Convert fragile Markdown constructs before HTML parsing.

    Standalone --- lines must become raw <hr /> elements; otherwise Markdown can
    reinterpret the preceding figure caption as a setext heading.
    """
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    out: list[str] = []
    display_math_blocks: list[str] = []
    inline_math_blocks: list[str] = []

    def replace_inline_math(line: str) -> str:
        def replacement(match: re.Match) -> str:
            formula = match.group(1) if match.group(1) is not None else match.group(2)
            placeholder = f"@@INLINE_MATH_{len(inline_math_blocks)}@@"
            inline_math_blocks.append(formula)
            return placeholder

        return re.sub(r"\\\((.+?)\\\)|\\\[(.+?)\\\]", replacement, line)

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped == "$$":
            i += 1
            block: list[str] = []
            while i < len(lines) and lines[i].strip() != "$$":
                block.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].strip() == "$$":
                i += 1
            placeholder = f"@@DISPLAY_MATH_{len(display_math_blocks)}@@"
            display_math_blocks.append("\n".join(block))
            out.append(placeholder)
            continue
        if stripped == "---":
            out.append("<hr />")
            i += 1
            continue
        out.append(replace_inline_math(lines[i]))
        i += 1
    return "\n".join(out), display_math_blocks, inline_math_blocks


def make_document(title: str | None = None) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (
        ("Title", 18),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
        ("Heading 4", 11),
    ):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        styles[style_name].font.size = Pt(11)

    if title:
        doc.core_properties.title = title
    doc.core_properties.author = "Codex"
    return doc


def resolve_image(src: str, markdown_path: Path) -> Path:
    src = unquote(src.strip())
    src = src.split("#", 1)[0].split("?", 1)[0]
    image_path = Path(src)
    if not image_path.is_absolute():
        image_path = (markdown_path.parent / image_path).resolve()
    return image_path


def add_image(document: Document, image_path: Path, alt: str = "") -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not image_path.exists():
        run = p.add_run(f"[Missing figure: {image_path}]")
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        return

    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
            dpi = img.info.get("dpi", (150, 150))[0] or 150
            width_in = width_px / dpi
            height_in = height_px / dpi
            target_width = min(CONTENT_WIDTH_IN, width_in)
            if target_width <= 0:
                target_width = CONTENT_WIDTH_IN
            target_height = height_in * (target_width / width_in) if width_in else None
    except Exception:
        target_width = CONTENT_WIDTH_IN
        target_height = None

    run = p.add_run()
    if target_height:
        run.add_picture(str(image_path), width=Inches(target_width), height=Inches(target_height))
    else:
        run.add_picture(str(image_path), width=Inches(target_width))
    if alt:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = caption.add_run(alt)
        r.italic = True
        r.font.size = Pt(9)


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


GREEK_COMMANDS = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "Delta": "Δ",
    "Pi": "Π",
    "Theta": "Θ",
    "varepsilon": "ε",
    "epsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "rho": "ρ",
    "sigma": "σ",
    "bar\\sigma": "σ̄",
    "tau": "τ",
    "varphi": "φ",
    "phi": "φ",
    "chi": "χ",
    "kappa": "κ",
    "psi": "ψ",
    "omega": "ω",
}

SYMBOL_COMMANDS = {
    "arg": "arg",
    "ast": "*",
    "Big": "",
    "big": "",
    "cdot": "·",
    "times": "×",
    "le": "≤",
    "leq": "≤",
    "ge": "≥",
    "geq": "≥",
    "neq": "≠",
    "in": "∈",
    "notin": "∉",
    "to": "→",
    "rightarrow": "→",
    "leftarrow": "←",
    "pm": "±",
    "approx": "≈",
    "equiv": "≡",
    "sim": "∼",
    "sum": "∑",
    "infty": "∞",
    "ell": "ℓ",
    "exp": "exp",
    "min": "min",
    "max": "max",
    "cos": "cos",
    "sin": "sin",
    "lVert": "‖",
    "rVert": "‖",
    "lVert": "‖",
    "rVert": "‖",
    "Vert": "‖",
    "langle": "⟨",
    "rangle": "⟩",
}

CAL_MAP = {
    "A": "𝒜",
    "C": "𝒞",
    "E": "ℰ",
    "F": "ℱ",
    "H": "ℋ",
    "K": "𝒦",
    "L": "ℒ",
    "N": "𝒩",
    "S": "𝒮",
}

BB_MAP = {
    "1": "𝟙",
    "N": "ℕ",
    "R": "ℝ",
    "Z": "ℤ",
}

BOLD_MAP = {
    "a": "𝐚",
    "b": "𝐛",
    "c": "𝐜",
    "d": "𝐝",
    "e": "𝐞",
    "f": "𝐟",
    "n": "𝐧",
    "t": "𝐭",
    "v": "𝐯",
    "x": "𝐱",
}


def matching_brace(text: str, open_index: int) -> int:
    depth = 0
    for idx in range(open_index, len(text)):
        if text[idx] == "{":
            depth += 1
        elif text[idx] == "}":
            depth -= 1
            if depth == 0:
                return idx
    return -1


def replace_command_one_arg(text: str, command: str, renderer) -> str:
    needle = f"\\{command}"
    pos = 0
    while True:
        idx = text.find(needle + "{", pos)
        if idx == -1:
            return text
        open_idx = idx + len(needle)
        close_idx = matching_brace(text, open_idx)
        if close_idx == -1:
            pos = idx + len(needle)
            continue
        arg = text[open_idx + 1 : close_idx]
        replacement = renderer(arg)
        text = text[:idx] + replacement + text[close_idx + 1 :]
        pos = idx + len(replacement)


def replace_frac(text: str) -> str:
    pos = 0
    while True:
        idx = text.find(r"\frac{", pos)
        if idx == -1:
            return text
        num_open = idx + len(r"\frac")
        num_close = matching_brace(text, num_open)
        if num_close == -1 or num_close + 1 >= len(text) or text[num_close + 1] != "{":
            pos = idx + 5
            continue
        den_open = num_close + 1
        den_close = matching_brace(text, den_open)
        if den_close == -1:
            pos = idx + 5
            continue
        numerator = text[num_open + 1 : num_close]
        denominator = text[den_open + 1 : den_close]
        replacement = f"({numerator})/({denominator})"
        text = text[:idx] + replacement + text[den_close + 1 :]
        pos = idx + len(replacement)


def replace_underbrace(text: str) -> str:
    pos = 0
    while True:
        idx = text.find(r"\underbrace{", pos)
        if idx == -1:
            return text
        arg_open = idx + len(r"\underbrace")
        arg_close = matching_brace(text, arg_open)
        if arg_close == -1:
            pos = idx + 11
            continue
        expression = text[arg_open + 1 : arg_close]
        end = arg_close + 1
        if text[end : end + 2] == "_{":
            label_open = end + 1
            label_close = matching_brace(text, label_open)
            if label_close != -1:
                end = label_close + 1
        text = text[:idx] + expression + text[end:]
        pos = idx + len(expression)


def normalize_latex_line(line: str) -> str:
    line = line.strip().rstrip(",")
    line = line.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
    line = line.replace("&", "")
    line = line.replace(r"\left", "").replace(r"\right", "")
    line = line.replace(r"\,", " ").replace(r"\;", " ").replace(r"\!", "")
    line = line.replace(r"\quad", "    ").replace(r"\qquad", "      ")
    line = line.replace(r"\ ", " ")
    line = line.replace(r"\{", "{").replace(r"\}", "}")
    line = line.replace(r"\[", "[").replace(r"\]", "]")
    line = line.replace(r"\|", "‖")

    line = replace_underbrace(line)
    line = replace_frac(line)
    line = replace_command_one_arg(line, "text", lambda arg: arg)
    line = replace_command_one_arg(line, "mathrm", lambda arg: arg)
    line = replace_command_one_arg(line, "operatorname", lambda arg: arg)
    line = replace_command_one_arg(line, "mathcal", lambda arg: CAL_MAP.get(arg, arg))
    line = replace_command_one_arg(line, "mathbb", lambda arg: BB_MAP.get(arg, arg))
    line = replace_command_one_arg(line, "mathbf", lambda arg: "".join(BOLD_MAP.get(ch, ch) for ch in arg))
    line = replace_command_one_arg(line, "boldsymbol", lambda arg: arg)
    line = replace_command_one_arg(line, "widehat", lambda arg: f"{arg}̂")
    line = replace_command_one_arg(line, "bar", lambda arg: f"{arg}̄")

    for command, symbol in sorted(GREEK_COMMANDS.items(), key=lambda kv: -len(kv[0])):
        line = line.replace(f"\\{command}", symbol)
    for command, symbol in sorted(SYMBOL_COMMANDS.items(), key=lambda kv: -len(kv[0])):
        line = line.replace(f"\\{command}", symbol)

    line = re.sub(r"\\([A-Za-z]+)", r"\1", line)
    line = line.replace("~", " ")
    line = re.sub(r"\s+", " ", line)
    line = re.sub(r"\s+([,.;:])", r"\1", line)
    return line.strip()


def split_latex_display(raw_text: str) -> tuple[list[str], str | None]:
    raw_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    tag_match = re.search(r"\\tag\{([^}]+)\}", raw_text)
    tag = tag_match.group(1).strip() if tag_match else None
    if tag_match:
        raw_text = raw_text[: tag_match.start()] + raw_text[tag_match.end() :]
    raw_text = raw_text.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
    raw_text = re.sub(r"\\\\\s*", "\n", raw_text)
    lines = [normalize_latex_line(line) for line in raw_text.splitlines()]
    return [line for line in lines if line], tag


def add_formula_runs(paragraph, formula: str) -> None:
    idx = 0
    normal_buffer: list[str] = []

    def flush_normal() -> None:
        if not normal_buffer:
            return
        run = paragraph.add_run("".join(normal_buffer))
        run.font.name = "Cambria Math"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
        run.font.size = Pt(11)
        normal_buffer.clear()

    while idx < len(formula):
        char = formula[idx]
        if char in "_^":
            flush_normal()
            script = "sub" if char == "_" else "sup"
            idx += 1
            if idx < len(formula) and formula[idx] == "{":
                end = matching_brace(formula, idx)
                if end == -1:
                    normal_buffer.append(char)
                    continue
                script_text = formula[idx + 1 : end]
                idx = end + 1
            elif idx < len(formula):
                script_text = formula[idx]
                idx += 1
            else:
                script_text = ""
            script_text = script_text.replace("{", "").replace("}", "")
            run = paragraph.add_run(script_text)
            run.font.name = "Cambria Math"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
            run.font.size = Pt(9)
            if script == "sub":
                run.font.subscript = True
            else:
                run.font.superscript = True
            continue
        if char in "{}":
            idx += 1
            continue
        normal_buffer.append(char)
        idx += 1
    flush_normal()


def m_el(tag: str):
    return OxmlElement(f"m:{tag}")


def omml_run(text: str):
    run = m_el("r")
    text_el = m_el("t")
    if text.startswith(" ") or text.endswith(" "):
        text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def omml_arg(tag: str, children):
    element = m_el(tag)
    for child in children:
        element.append(child)
    return element


def omml_script(base, sub=None, sup=None):
    if sub is not None and sup is not None:
        element = m_el("sSubSup")
        element.append(omml_arg("e", base))
        element.append(omml_arg("sub", sub))
        element.append(omml_arg("sup", sup))
        return element
    if sub is not None:
        element = m_el("sSub")
        element.append(omml_arg("e", base))
        element.append(omml_arg("sub", sub))
        return element
    if sup is not None:
        element = m_el("sSup")
        element.append(omml_arg("e", base))
        element.append(omml_arg("sup", sup))
        return element
    return base[0] if len(base) == 1 else omml_arg("e", base)


def omml_fraction(numerator, denominator):
    element = m_el("f")
    element.append(omml_arg("num", numerator))
    element.append(omml_arg("den", denominator))
    return element


def omml_radical(radicand):
    element = m_el("rad")
    element.append(omml_arg("deg", []))
    element.append(omml_arg("e", radicand))
    return element


def prepare_latex_for_omml(line: str) -> str:
    line = line.strip().rstrip(",")
    line = line.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
    line = line.replace("&", "")
    line = line.replace(r"\left", "").replace(r"\right", "")
    line = line.replace(r"\,", " ").replace(r"\;", " ").replace(r"\!", "")
    line = line.replace(r"\qquad", "      ").replace(r"\quad", "    ")
    line = line.replace(r"\ ", " ")
    line = line.replace(r"\{", "{").replace(r"\}", "}")
    line = line.replace(r"\[", "[").replace(r"\]", "]")
    return line


class OmmlLatexParser:
    def __init__(self, text: str):
        self.text = prepare_latex_for_omml(text)
        self.index = 0

    def current(self) -> str:
        if self.index >= len(self.text):
            return ""
        return self.text[self.index]

    def startswith(self, value: str) -> bool:
        return self.text.startswith(value, self.index)

    def parse(self, stop: str | None = None):
        elements = []
        while self.index < len(self.text):
            if stop and self.current() == stop:
                break
            if self.current() == "}":
                break
            atom = self.parse_atom_with_scripts()
            if atom:
                elements.extend(atom if isinstance(atom, list) else [atom])
        return elements

    def parse_atom_with_scripts(self):
        base = self.parse_atom()
        if not base:
            return []

        sub = None
        sup = None
        while self.current() in {"_", "^"}:
            marker = self.current()
            self.index += 1
            arg = self.parse_script_argument()
            if marker == "_":
                sub = arg
            else:
                sup = arg

        if sub is not None or sup is not None:
            return [omml_script(base, sub=sub, sup=sup)]
        return base

    def parse_script_argument(self):
        if self.current() == "{":
            self.index += 1
            value = self.parse(stop="}")
            if self.current() == "}":
                self.index += 1
            return value
        return self.parse_atom()

    def parse_braced_argument(self):
        while self.current().isspace():
            self.index += 1
        if self.current() != "{":
            return []
        self.index += 1
        value = self.parse(stop="}")
        if self.current() == "}":
            self.index += 1
        return value

    def parse_raw_braced_text(self) -> str:
        while self.current().isspace():
            self.index += 1
        if self.current() != "{":
            return ""
        start = self.index
        end = matching_brace(self.text, start)
        if end == -1:
            return ""
        raw = self.text[start + 1 : end]
        self.index = end + 1
        return raw

    def parse_atom(self):
        if self.index >= len(self.text):
            return []

        ch = self.current()
        if ch.isspace():
            start = self.index
            while self.index < len(self.text) and self.text[self.index].isspace():
                self.index += 1
            return [omml_run(" " if self.index > start else "")]

        if ch == "{":
            self.index += 1
            group = self.parse(stop="}")
            if self.current() == "}":
                self.index += 1
            return group

        if ch == "\\":
            return self.parse_command()

        start = self.index
        while self.index < len(self.text):
            ch = self.text[self.index]
            if ch in r"\{}_^":
                break
            if ch.isspace():
                break
            self.index += 1
        if self.index == start:
            self.index += 1
            return [omml_run(ch)]
        return [omml_run(self.text[start : self.index])]

    def parse_command(self):
        self.index += 1
        start = self.index
        while self.index < len(self.text) and self.text[self.index].isalpha():
            self.index += 1
        command = self.text[start : self.index]
        if not command and self.index < len(self.text):
            char = self.text[self.index]
            self.index += 1
            return [omml_run(char)]

        if command in {"frac", "tfrac"}:
            numerator = self.parse_braced_argument()
            denominator = self.parse_braced_argument()
            return [omml_fraction(numerator, denominator)]
        if command == "sqrt":
            radicand = self.parse_braced_argument()
            return [omml_radical(radicand)]
        if command == "underbrace":
            expression = self.parse_braced_argument()
            if self.current() == "_":
                self.index += 1
                _ = self.parse_script_argument()
            return expression
        if command in {"text", "mathrm", "operatorname"}:
            raw = self.parse_raw_braced_text()
            return [omml_run(raw)]
        if command == "mathcal":
            raw = self.parse_raw_braced_text()
            return [omml_run(CAL_MAP.get(raw, raw))]
        if command == "mathbb":
            raw = self.parse_raw_braced_text()
            return [omml_run(BB_MAP.get(raw, raw))]
        if command == "mathbf":
            raw = self.parse_raw_braced_text()
            return [omml_run("".join(BOLD_MAP.get(char, char) for char in raw))]
        if command == "boldsymbol":
            return self.parse_braced_argument()
        if command in {"widehat", "hat"}:
            raw = self.parse_raw_braced_text()
            return [omml_run(f"{normalize_latex_line(raw)}̂")]
        if command == "bar":
            raw = self.parse_raw_braced_text()
            return [omml_run(f"{normalize_latex_line(raw)}̄")]
        if command in GREEK_COMMANDS:
            return [omml_run(GREEK_COMMANDS[command])]
        if command in SYMBOL_COMMANDS:
            symbol = SYMBOL_COMMANDS[command]
            return [omml_run(symbol)] if symbol else []
        return [omml_run(command)]


def latex_line_to_omml(line: str):
    return OmmlLatexParser(line).parse()


def add_plain_text_run(paragraph, text: str, bold=False, italic=False, code=False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9.5)


def add_inline_omml(paragraph, formula: str) -> None:
    omath = m_el("oMath")
    for child in latex_line_to_omml(formula):
        omath.append(child)
    paragraph._p.append(omath)


def add_text_with_inline_math(paragraph, text: str, bold=False, italic=False, code=False) -> None:
    if code:
        add_plain_text_run(paragraph, text, bold=bold, italic=italic, code=True)
        return

    pattern = re.compile(r"@@INLINE_MATH_(\d+)@@|\\\((.+?)\\\)|\\\[(.+?)\\\]")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            add_plain_text_run(paragraph, text[pos : match.start()], bold=bold, italic=italic)
        if match.group(1) is not None:
            idx = int(match.group(1))
            formula = CURRENT_INLINE_MATH_BLOCKS[idx] if 0 <= idx < len(CURRENT_INLINE_MATH_BLOCKS) else ""
        else:
            formula = match.group(2) if match.group(2) is not None else match.group(3)
        add_inline_omml(paragraph, formula)
        pos = match.end()
    if pos < len(text):
        add_plain_text_run(paragraph, text[pos:], bold=bold, italic=italic)


def add_runs(paragraph, node, bold=False, italic=False, code=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            add_text_with_inline_math(paragraph, text, bold=bold, italic=italic, code=code)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in {"strong", "b"}:
        for child in node.children:
            add_runs(paragraph, child, bold=True or bold, italic=italic, code=code)
    elif name in {"em", "i"}:
        for child in node.children:
            add_runs(paragraph, child, bold=bold, italic=True or italic, code=code)
    elif name == "code":
        for child in node.children:
            add_runs(paragraph, child, bold=bold, italic=italic, code=True)
    elif name == "br":
        paragraph.add_run().add_break()
    elif name == "a":
        text = clean_text(node.get_text(" "))
        href = node.get("href")
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(5, 99, 193)
        run.underline = True
        if href:
            paragraph.add_run(f" ({href})")
    elif name == "img":
        paragraph.add_run(f"[Figure: {node.get('alt', node.get('src', ''))}]")
    elif name == "sup":
        run = paragraph.add_run(clean_text(node.get_text(" ")))
        run.font.superscript = True
    elif name == "sub":
        run = paragraph.add_run(clean_text(node.get_text(" ")))
        run.font.subscript = True
    else:
        for child in node.children:
            add_runs(paragraph, child, bold=bold, italic=italic, code=code)


def add_paragraph_from_tag(document: Document, tag: Tag, style: str | None = None) -> None:
    text = clean_text(tag.get_text(" "))
    if not text and not tag.find("img"):
        return
    p = document.add_paragraph(style=style)
    for child in tag.children:
        add_runs(p, child)


def add_list(document: Document, tag: Tag, ordered: bool = False, level: int = 0) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in tag.find_all("li", recursive=False):
        p = document.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0.6 + level * 0.45)
        inline_children = [child for child in li.children if not (isinstance(child, Tag) and child.name in {"ul", "ol"})]
        for child in inline_children:
            add_runs(p, child)
        for child in li.children:
            if isinstance(child, Tag) and child.name == "ul":
                add_list(document, child, ordered=False, level=level + 1)
            elif isinstance(child, Tag) and child.name == "ol":
                add_list(document, child, ordered=True, level=level + 1)


def add_table(document: Document, table_tag: Tag) -> None:
    rows = []
    for tr in table_tag.find_all("tr"):
        cells = tr.find_all(["th", "td"], recursive=False)
        rows.append(cells)
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    cell_width = int(9026 / max_cols)

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_width(cell, cell_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if r_idx == 0:
                set_cell_shading(cell, "EDEDED")
            text_container = row[c_idx] if c_idx < len(row) else None
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if text_container is not None:
                for child in text_container.children:
                    add_runs(paragraph, child, bold=(r_idx == 0))
            for run in paragraph.runs:
                run.font.size = Pt(8.5 if max_cols >= 5 else 9.5)

    document.add_paragraph()


def add_display_math_text(document: Document, raw_text: str) -> None:
    tag_match = re.search(r"\\tag\{([^}]+)\}", raw_text)
    tag_number = tag_match.group(1).strip() if tag_match else None
    if tag_match:
        raw_text = raw_text[: tag_match.start()] + raw_text[tag_match.end() :]
    raw_text = raw_text.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
    raw_lines = [line.strip() for line in re.split(r"\\\\\s*|\n", raw_text) if line.strip()]
    if not raw_lines:
        return

    for idx, line in enumerate(raw_lines):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(4 if idx == len(raw_lines) - 1 else 0)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(CONTENT_WIDTH_IN / 2), WD_TAB_ALIGNMENT.CENTER)
        tabs.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        omath = m_el("oMath")
        for child in latex_line_to_omml(line):
            omath.append(child)
        p._p.append(omath)
        if tag_number and idx == len(raw_lines) - 1:
            number_run = p.add_run(f"\t({tag_number})")
            number_run.font.name = "Times New Roman"
            number_run.font.size = Pt(11)


def add_display_math(document: Document, tag: Tag) -> None:
    add_display_math_text(document, "\n".join(tag.stripped_strings))


def add_code_block(document: Document, tag: Tag) -> None:
    text = tag.get_text("\n").rstrip()
    if not text:
        return
    for line in text.splitlines():
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8.5)


def render_node(document: Document, node, markdown_path: Path, display_math_blocks: list[str]) -> None:
    if isinstance(node, NavigableString):
        if str(node).strip():
            document.add_paragraph(str(node).strip())
        return
    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(name[1]), 4)
        text = clean_text(node.get_text(" "))
        if level == 1 and len(document.paragraphs) <= 1:
            p = document.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(text)
        else:
            document.add_heading(text, level=level)
    elif name == "p":
        placeholder_text = clean_text(node.get_text(" "))
        placeholder_match = re.fullmatch(r"@@DISPLAY_MATH_(\d+)@@", placeholder_text)
        if placeholder_match:
            block_index = int(placeholder_match.group(1))
            if 0 <= block_index < len(display_math_blocks):
                add_display_math_text(document, display_math_blocks[block_index])
            return
        images = node.find_all("img", recursive=False)
        non_image_text = clean_text(" ".join(child.get_text(" ") if isinstance(child, Tag) and child.name != "img" else str(child) for child in node.children if not (isinstance(child, Tag) and child.name == "img")))
        if images and not non_image_text:
            for img in images:
                add_image(document, resolve_image(img.get("src", ""), markdown_path), img.get("alt", ""))
        else:
            add_paragraph_from_tag(document, node)
    elif name == "blockquote":
        for child in node.children:
            if isinstance(child, Tag) and child.name == "p":
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.right_indent = Cm(0.4)
                for inner in child.children:
                    add_runs(p, inner, italic=True)
    elif name == "ul":
        add_list(document, node, ordered=False)
    elif name == "ol":
        add_list(document, node, ordered=True)
    elif name == "table":
        add_table(document, node)
    elif name == "pre":
        add_code_block(document, node)
    elif name == "hr":
        add_horizontal_rule(document)
    elif name == "div" and "display-math" in node.get("class", []):
        add_display_math(document, node)
    else:
        for child in node.children:
            render_node(document, child, markdown_path, display_math_blocks)


def add_footer(document: Document, footer_text: str) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(footer_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(96, 96, 96)
        run.add_text(" | Page ")
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)


def convert_markdown_to_docx(input_path: Path, output_path: Path, footer_text: str) -> None:
    global CURRENT_INLINE_MATH_BLOCKS
    raw = input_path.read_text(encoding="utf-8")
    preprocessed, display_math_blocks, inline_math_blocks = preprocess_display_math(raw)
    CURRENT_INLINE_MATH_BLOCKS = inline_math_blocks
    html_text = markdown.markdown(
        preprocessed,
        extensions=["extra", "tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html_text, "html.parser")
    first_heading = soup.find(re.compile("^h[1-6]$"))
    title = clean_text(first_heading.get_text(" ")) if first_heading else input_path.stem
    doc = make_document(title)

    for child in soup.contents:
        render_node(doc, child, input_path, display_math_blocks)
    add_footer(doc, footer_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--footer", default="IJDRR draft for supervisor review")
    args = parser.parse_args()
    convert_markdown_to_docx(args.input, args.output, args.footer)


if __name__ == "__main__":
    main()
